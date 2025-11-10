import os
import sys
import tempfile
from pathlib import Path

import pytest

try:
    from docx import Document
except Exception:
    Document = None

from text_file_editing import docx_to_md


@pytest.mark.skipif(Document is None, reason="python-docx not installed")
def test_convert_simple_docx(tmp_path: Path):
    # Create a simple .docx fixture
    doc_path = tmp_path / 'simple.docx'
    doc = Document()
    doc.add_heading('Title', level=1)
    doc.add_paragraph('This is a paragraph.')
    doc.add_paragraph('• Bullet item')
    doc.save(str(doc_path))

    out_path = tmp_path / 'simple.md'

    # Run conversion
    md = docx_to_md.convert_docx_to_md(str(doc_path), str(out_path))

    assert out_path.exists()
    text = out_path.read_text(encoding='utf-8')

    # Basic checks for conversion
    assert '# Title' in text
    assert 'This is a paragraph.' in text
    assert '- Bullet item' in text


def test_cli_gui_selector_no_tk(monkeypatch):
    # If tkinter is not available, choose_file_via_tk should raise RuntimeError
    # Correctly remove tkinter from sys.modules using the actual mapping
    monkeypatch.setitem(sys.modules, 'tkinter', None)
    with pytest.raises(RuntimeError):
        docx_to_md.choose_file_via_tk()
